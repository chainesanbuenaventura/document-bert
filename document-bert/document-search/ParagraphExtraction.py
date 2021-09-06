import os
import re
import pandas as pd
import glob
import numpy as np
import matplotlib.pyplot as plt

from PyPDF2 import PdfFileWriter, PdfFileReader
from docx import Document
from tkinter import Tcl
from urllib.parse import urlparse
from pathlib import Path


class ParagraphExtractor: 
    def __init__(self, file_path):
        self.file_path = file_path
        self.document = Document(file_path)
        self.curr_header = ''
        self.curr_subheader = ''
        self.appendix_found = False
        
        self.paragraphs = []
        self.main_headers = []
        self.sub_headers = []
        
    def is_header(self, paragraph_obj):
        try:
            if (paragraph_obj.style.font.bold 
                and paragraph_obj.style.font.size.pt == 11 
                and paragraph_obj.text.isupper()):
                return True 
        except:
            return False 
        return False 
    
    def is_subheader(self, paragraph_obj):
        try:
            if ((paragraph_obj.style.name == 'Title'
                or paragraph_obj.style.name == 'Heading 1'
                or paragraph_obj.style.name == 'List Paragraph')
                and not paragraph_obj.text.isupper()
                and len(paragraph_obj.text.split(' ')) < 10):
                    return True
        except:
            return False 
        return False
    
    def is_paragraph_ok(self, paragraph_obj):
        try:
            #paragraph should not be empty
            if len(paragraph_obj.text.split()) <= 10:
                return False
            #paragrah should not be a footer
#             elif re.search("^\d+\s(.+)", paragraph_obj.text): 
#                 return False
#             #paragraph should not have a normal style
#             elif paragraph_obj.style.name == 'Normal':
#                 return False
#             #paragraph should have a font size of 11 
#             elif paragraph_obj.style.font.size.pt < 11:
#                 return False 
        except:
            #paragraph should have more than 10 words 
            if len(paragraph_obj.text.split(' ')) < 10:
                return False
        return True 
        
    def is_main_paragraph(self, paragraph_obj):
        if (paragraph_obj.style.name == 'List Paragraph'
            or (re.search("^\d+\.\s(.+)",paragraph_obj.text)
                and len(paragraph_obj.text.split(' ')) > 50)
            or len(self.paragraphs) == 0):
                return True
        return False
    
    def extract_paragraphs(self):
        print(f"Extracting {self.file_path}")
        paragraphs = []
        for paragraph_obj in self.document.paragraphs: 
            
            if self.is_paragraph_ok(paragraph_obj):
                paragraphs.append(paragraph_obj.text)
                
        return paragraphs